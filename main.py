import os
from dotenv import load_dotenv
import subprocess
import asyncio
from contextlib import asynccontextmanager
from typing import Optional
from agent import ChatAgent  
from fastapi import FastAPI, Request, Header, HTTPException, BackgroundTasks
from fastapi.concurrency import run_in_threadpool
from httpx import AsyncClient
from pydantic import BaseModel
from docx import Document
from docx.shared import Pt
from mcp_server import get_server_port

# Load environment variables from .env file
load_dotenv()

# Initialize the agent globally (outside the function) so it persists
agent = ChatAgent()


# Create FastAPI app with lifespan
app = FastAPI()

BOT_TOKEN = os.getenv("BOT_TOKEN")
SECRET_TOKEN = os.getenv("SECRET_TOKEN")
MY_CHAT_ID = os.getenv("MY_CHAT_ID")
TELEGRAM_API_URL = f"https://api.telegram.org/bot{BOT_TOKEN}"
KNOWLEDGE_BASE_PATH = "knowledge_base.docx"


# Pydantic model for feed request body
class FeedRequest(BaseModel):
    text: str


def append_to_word_doc(text: str) -> None:
    """
    Append text to the Word document.

    Args:
        text: The text to append to the document
    """
    try:
        # Create document if it doesn't exist
        if os.path.exists(KNOWLEDGE_BASE_PATH):
            doc = Document(KNOWLEDGE_BASE_PATH)
        else:
            doc = Document()
            # Add a title for new documents
            title = doc.add_heading("Knowledge Base", 0)
            title.runs[0].font.size = Pt(14)

        # Append the new text as a paragraph
        doc.add_paragraph(text)

        # Save the document
        doc.save(KNOWLEDGE_BASE_PATH)
    except PermissionError as e:
        print(f"Permission error when writing to {KNOWLEDGE_BASE_PATH}: {e}")
        raise
    except Exception as e:
        print(f"Error writing to {KNOWLEDGE_BASE_PATH}: {e}")
        raise



@app.get("/health")
async def health():
    """Health check endpoint for UptimeRobot to keep the server awake."""
    return {"status": "ok", "mcp_server": "running" if mcp_server_process else "stopped"}


@app.get("/")
async def root():
    """Root endpoint."""
    return {
        "message": "Telegram RAG Summarizer API is running",      
        "endpoints": {
            "health": "/health",
            "admin_feed": "/admin/feed",          
            "webhook": "/webhook"
        }
    }


async def send_message(chat_id: int, text: str):
    """Send a message to the specified chat ID via Telegram API."""
    async with AsyncClient() as client:
        url = f"{TELEGRAM_API_URL}/sendMessage"
        payload = {
            "chat_id": chat_id,
            "text": text,
        }
        response = await client.post(url, json=payload)
        response.raise_for_status()


@app.post("/webhook")
async def webhook(
    request: Request,
    x_telegram_bot_api_secret_token: Optional[str] = Header(None)
):
    """
    Webhook endpoint to receive messages from Telegram.
    Invokes the ChatAgent to perform RAG search and summarization.
    """
   
    print(x_telegram_bot_api_secret_token)
    # 1. Verify webhook secret token (Security)
    if x_telegram_bot_api_secret_token != SECRET_TOKEN:
        raise HTTPException(status_code=403, detail="Unauthorized webhook")

    # 2. Parse the incoming update
    update = await request.json()
    message = update.get("message", {})
    chat_id = message.get("chat", {}).get("id")
    from_user_id = str(message.get("from", {}).get("id"))

    # 3. Authorization Check
    #if from_user_id != MY_CHAT_ID:
    #    return {"status": "ignored", "reason": "unauthorized_user"}

    # 4. Extract text and invoke the Agent
    text = message.get("text")
    if text:
        try:
            # We use 'await' because ChatAgent.run is an async function
            # This triggers the search_docs tool and the Qwen summarization
            print("before invking agent")
            ai_response = await agent.run(text)
            
            # 5. Send the final generated response back to Telegram
            await send_message(chat_id, ai_response)
            
        except Exception as e:
            # Log the error and notify the user so they aren't left waiting
            print(f"Error in Agent Execution: {e}")
            await send_message(chat_id, "I'm sorry, I had trouble accessing my knowledge base. Please try again in a moment.")

    return {"status": "ok"}

@app.post("/admin/feed")
async def feed_knowledge_base(request: FeedRequest, background_tasks: BackgroundTasks):
    """
    Administrative endpoint to ingest new text into the knowledge base.

    The text is appended to knowledge_base.docx and the MCP server index
    is refreshed in the background to make the new data searchable.

    Args:
        request: FeedRequest containing the text to add
        background_tasks: FastAPI BackgroundTasks for non-blocking operations

    Returns:
        JSON response confirming the operation
    """
    if not request.text or not request.text.strip():
        raise HTTPException(status_code=400, detail="Text cannot be empty")

    try:
        # Run the document write operation in a threadpool to prevent blocking the event loop
        await run_in_threadpool(append_to_word_doc, request.text)
    except PermissionError as e:
        raise HTTPException(
            status_code=403,
            detail=f"Permission denied: Unable to write to knowledge base. {str(e)}"
        )
    except FileNotFoundError as e:
        raise HTTPException(
            status_code=500,
            detail=f"File system error: {str(e)}"
        )
    except OSError as e:
        raise HTTPException(
            status_code=500,
            detail=f"File system error: {str(e)}"
        )
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to write to knowledge base: {str(e)}"
        )

    # Add the MCP index refresh as a background task
    background_tasks.add_task(refresh_mcp_index)

    return {
        "success": True,
        "message": f"Text successfully added to knowledge base ({len(request.text)} characters)",
        "reindexing": "in_progress"
    }





